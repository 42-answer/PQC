#!/usr/bin/env python3
"""Read the Word template and extract its structure"""

from docx import Document

doc = Document('Sample_Template_For_Solution_Submission.docx')

print("=" * 60)
print("TEMPLATE STRUCTURE")
print("=" * 60)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"\n[{i}] {text}")

print("\n" + "=" * 60)
print("TABLES FOUND")
print("=" * 60)

for i, table in enumerate(doc.tables):
    print(f"\nTable {i+1}:")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))
